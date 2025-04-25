import docx

def sub_top_g10(file,opt=''):
    global ds, bd, x1, g ,doc1, doc
    f = open(file)
    if opt == '':
        doc = docx.Document()
    else:
        namelist = open(opt)
        doc = docx.Document()
        doc1 = docx.Document()
        x1 = namelist.readlines()
    x = f.readlines()
    lst = []
    l2 = []
    final = []

    def sorting(sub_name):
        for s in range(len(sub_name)):
            for total_lst in range(0, len(sub_name) - s - 1):
                if sub_name[total_lst][-1] < sub_name[total_lst + 1][-1]:
                    sub_name[total_lst], sub_name[total_lst + 1] = sub_name[total_lst + 1], sub_name[total_lst]


    for i in range(len(x)):
        y = x[i].split()
        if y == []:
            pass
        else:
            roll = ''
            name = ''
            marks_lst = []

            for j in y:
                name_rollno = [roll, name[1:]]
                if y[0].isdigit() and j.isalpha() and j != ('PASS' or 'FAIL'):
                    name = name + j + ' '
                elif y[0].isdigit() and j.isdigit() and len(j) > 4:
                    roll += j

            if len(name_rollno[0]) > 1 and len(name_rollno[1]) > 1:
                lst.append(name_rollno)

            if y[0].isdigit() and len(y[0]) > 4:
                sub_codes = []
                for code in y:
                    if code.isdigit() and len(code) == 3:
                        sub_codes.append(code)

                if sub_codes[1] == '006':
                    y1 = x[i + 1].split()
                    for k in range(len(y1)):
                        if k % 2 == 0:
                            marks_lst.append(y1[k])
                    marks_lst.insert(2, '-')
                    marks_lst.insert(3, '-')
                elif sub_codes[1] == '085':
                    y1 = x[i + 1].split()
                    for k in range(len(y1)):
                        if k % 2 == 0:
                            marks_lst.append(y1[k])
                    marks_lst.insert(1, '-')
                    marks_lst.insert(3, '-')
                elif sub_codes[1] == '018':
                    y1 = x[i + 1].split()
                    for k in range(len(y1)):
                        if k % 2 == 0:
                            marks_lst.append(y1[k])
                    marks_lst.insert(1, '-')
                    marks_lst.insert(2, '-')

            if marks_lst == []:
                pass
            else:
                l2.append(marks_lst)


    for j in range(0, len(lst)):
        lk = lst[j] + l2[j]
        final.append(lk)

    if opt == '':
        eng = []
        mat = []
        tam = []
        hin = []
        frc = []
        sci = []
        soc = []
        for a in final:
            if a[2].isdigit():
                eng.append(a[1:3])
            if a[3].isdigit():
                tam.append([a[1]] + [a[3]])
            if a[4].isdigit():
                hin.append([a[1]] + [a[4]])
            if a[5].isdigit():
                frc.append([a[1]] + [a[5]])
            if a[6].isdigit():
                mat.append([a[1]] + [a[6]])
            if a[7].isdigit():
                sci.append([a[1]] + [a[7]])
            if a[8].isdigit():
                soc.append([a[1]] + [a[8]])

        sorting(eng);
        sorting(tam);
        sorting(hin);
        sorting(frc);
        sorting(mat);
        sorting(sci);
        sorting(soc)

        count_eng = 1
        for t in range(len(eng)):
            if count_eng == 4:
                eng = eng[:t]
                break
            elif eng[t][-1] != eng[t + 1][-1]:
                eng[t].insert(0, count_eng)
                count_eng += 1
            else:
                eng[t].insert(0, count_eng)

        count_tam = 1
        for t in range(len(tam)):
            if count_tam == 4:
                tam = tam[:t]
                break
            elif tam[t][-1] != tam[t + 1][-1]:
                tam[t].insert(0, count_tam)
                count_tam += 1
            else:
                tam[t].insert(0, count_tam)

        count_hin = 1
        for t in range(len(hin)):
            if count_hin == 4:
                hin = hin[:t]
                break
            elif hin[t][-1] != hin[t + 1][-1]:
                hin[t].insert(0, count_hin)
                count_hin += 1
            else:
                hin[t].insert(0, count_hin)

        count_frc = 1
        for t in range(len(frc)):
            if count_frc == 4:
                frc = frc[:t]
                break
            elif frc[t][-1] != frc[t + 1][-1]:
                frc[t].insert(0, count_frc)
                count_frc += 1
            else:
                frc[t].insert(0, count_frc)

        count_mat = 1
        for t in range(len(mat)):
            if count_mat == 4:
                mat = mat[:t]
                break
            elif mat[t][-1] != mat[t + 1][-1]:
                mat[t].insert(0, count_mat)
                count_mat += 1
            else:
                mat[t].insert(0, count_mat)

        count_sci = 1
        for t in range(len(sci)):
            if count_sci == 4:
                sci = sci[:t]
                break
            elif sci[t][-1] != sci[t + 1][-1]:
                sci[t].insert(0, count_sci)
                count_sci += 1
            else:
                sci[t].insert(0, count_sci)

        count_soc = 1
        for t in range(len(soc)):
            if count_soc == 4:
                soc = soc[:t]
                break
            elif soc[t][-1] != soc[t + 1][-1]:
                soc[t].insert(0, count_soc)
                count_soc += 1
            else:
                soc[t].insert(0, count_soc)

        def table(sub):
            if sub == eng:
                doc.add_heading("English Toppers", 0)
            if sub == mat:
                doc.add_heading("Maths Toppers", 0)
            if sub == tam:
                doc.add_heading("Tamil Toppers", 0)
            if sub == frc:
                doc.add_heading("French Toppers", 0)
            if sub == hin:
                doc.add_heading("Hindi Toppers", 0)
            if sub == sci:
                doc.add_heading("Science Toppers", 0)
            if sub == soc:
                doc.add_heading("Social Toppers", 0)
            table = doc.add_table(rows=1, cols=3)
            row = table.rows[0].cells
            row[0].text = 'Place'
            row[1].text = 'Name'
            row[2].text = 'Mark'
            for place, name, mark in sub:
                row = table.add_row().cells
                row[0].text = str(place)
                row[1].text = name
                row[2].text = str(int(mark))
            doc.add_page_break()
            doc.save(file[:-4] + '.docx')

        table(eng);
        table(tam);
        table(hin);
        table(frc);
        table(mat);
        table(sci);
        table(soc)

    else:
        name_lst = []
        for i in x1:
            name_lst.append(i.strip())

        ds_lst = []
        bd_lst = []
        for i in final:
            if i[0] not in name_lst:
                ds_lst.append(i)
            elif i[0] in name_lst:
                bd_lst.append(i)

        def option(var,f_ext):
            eng = []
            mat = []
            tam = []
            hin = []
            frc = []
            sci = []
            soc = []
            for a in var:
                if a[2].isdigit():
                    eng.append(a[1:3])
                if a[3].isdigit():
                    tam.append([a[1]] + [a[3]])
                if a[4].isdigit():
                    hin.append([a[1]] + [a[4]])
                if a[5].isdigit():
                    frc.append([a[1]] + [a[5]])
                if a[6].isdigit():
                    mat.append([a[1]] + [a[6]])
                if a[7].isdigit():
                    sci.append([a[1]] + [a[7]])
                if a[8].isdigit():
                    soc.append([a[1]] + [a[8]])

            sorting(eng);
            sorting(tam);
            sorting(hin);
            sorting(frc);
            sorting(mat);
            sorting(sci);
            sorting(soc)

            count_eng = 1
            for t in range(len(eng)):
                if count_eng == 4:
                    eng = eng[:t]
                    break
                elif eng[t][-1] != eng[t + 1][-1]:
                    eng[t].insert(0, count_eng)
                    count_eng += 1
                else:
                    eng[t].insert(0, count_eng)

            count_tam = 1
            for t in range(len(tam)):
                if count_tam == 4:
                    tam = tam[:t]
                    break
                elif tam[t][-1] != tam[t + 1][-1]:
                    tam[t].insert(0, count_tam)
                    count_tam += 1
                else:
                    tam[t].insert(0, count_tam)

            count_hin = 1
            if len(hin)!=0:
                for t in range(len(hin)):
                    if t==len(hin)-1:
                        if hin[t][-1] != hin[t -1][-1]:
                            hin[t].insert(0,count_hin)
                        elif count_hin!=1:
                            hin[t].insert(0,count_hin-1)
                        else:
                            hin[t].insert(0,count_hin)
                    elif count_hin == 4:
                        hin = hin[:t]
                        break
                    elif hin[t][-1] != hin[t + 1][-1]:
                        hin[t].insert(0, count_hin)
                        count_hin += 1
                    else:
                        hin[t].insert(0, count_hin)
            else:
                pass

            count_frc = 1
            if len(frc)!=0:
                for t in range(len(frc)):
                    if t==len(frc)-1:
                        if frc[t][-1] != frc[t -1][-1]:
                            frc[t].insert(0,count_frc)
                        elif count_frc!=1:
                            frc[t].insert(0,count_frc-1)
                        else:
                            frc[t].insert(0,count_frc)
                    elif count_frc == 4:
                        frc = frc[:t]
                        break
                    elif frc[t][-1] != frc[t + 1][-1]:
                        frc[t].insert(0, count_frc)
                        count_frc += 1
                    else:
                        frc[t].insert(0, count_frc)
            else:
                pass

            count_mat = 1
            for t in range(len(mat)):
                if count_mat == 4:
                    mat = mat[:t]
                    break
                elif mat[t][-1] != mat[t + 1][-1]:
                    mat[t].insert(0, count_mat)
                    count_mat += 1
                else:
                    mat[t].insert(0, count_mat)

            count_sci = 1
            for t in range(len(sci)):
                if count_sci == 4:
                    sci = sci[:t]
                    break
                elif sci[t][-1] != sci[t + 1][-1]:
                    sci[t].insert(0, count_sci)
                    count_sci += 1
                else:
                    sci[t].insert(0, count_sci)

            count_soc = 1
            for t in range(len(soc)):
                if count_soc == 4:
                    soc = soc[:t]
                    break
                elif soc[t][-1] != soc[t + 1][-1]:
                    soc[t].insert(0, count_soc)
                    count_soc += 1
                else:
                    soc[t].insert(0, count_soc)

            def table(sub,f_ext):
                if sub == eng:
                    f_ext.add_heading("English Toppers", 0)
                if sub == mat:
                    f_ext.add_heading("Maths Toppers", 0)
                if sub == tam:
                    f_ext.add_heading("Tamil Toppers", 0)
                if sub == frc :
                    f_ext.add_heading("French Toppers", 0)
                if sub == hin:
                    f_ext.add_heading("Hindi Toppers", 0)
                if sub == sci:
                    f_ext.add_heading("Science Toppers", 0)
                if sub == soc:
                    f_ext.add_heading("Social Toppers", 0)

                table = f_ext.add_table(rows=1, cols=3)
                row = table.rows[0].cells
                row[0].text = 'Place'
                row[1].text = 'Name'
                row[2].text = 'Mark'
                for place, name, mark in sub:
                    row = table.add_row().cells
                    row[0].text = str(place)
                    row[1].text = name
                    row[2].text = str(int(mark))
                    f_ext.add_page_break()
                    if f_ext == doc:
                        f_ext.save(file[:-4] + '(DS).docx')
                    else:
                        f_ext.save(file[:-4] + '(BD).docx')

            if len(eng) != 0:
                table(eng,f_ext)
            if len(tam) != 0:
                table(tam,f_ext)
            if len(hin)!=0:
                table(hin,f_ext)
            if len(frc)!=0:
                table(frc,f_ext)
            if len(mat) != 0:
                table(mat,f_ext)
            if len(sci) != 0:
                table(sci,f_ext)
            if len(soc) != 0:
                table(soc,f_ext)

            eng.clear(); mat.clear(); tam.clear(); hin.clear(); frc.clear(); sci.clear(); soc.clear()
        option(ds_lst, doc)
        option(bd_lst, doc1)

def sub_top_g12(file,opt=''):
    global ds, bd, x1, g, doc1 ,doc
    f = open(file)
    if opt == '':
        doc = docx.Document()
    else:
        namelist = open(opt)
        doc = docx.Document()
        doc1 = docx.Document()
        x1 = namelist.readlines()
    x = f.readlines()
    lst = []
    l2 = []
    eng = []
    mat = []
    physics = []
    chemistry = []
    C_S = []
    ap_maths = []
    biology = []
    economics = []
    account = []
    b_s = []
    final = []

    def sorting(sub_name):
        for s in range(len(sub_name)):
            for total_lst in range(0, len(sub_name) - s - 1):
                if sub_name[total_lst][-1] < sub_name[total_lst + 1][-1]:
                    sub_name[total_lst], sub_name[total_lst + 1] = sub_name[total_lst + 1], sub_name[total_lst]

    for i in range(len(x)):
        y = x[i].split()
        if y == []:
            pass
        else:
            roll = ''
            name = ''
            english, maths, phy, che, cs, ap_mat, bio, eco, acc, bs = '-', '-', '-', '-', '-', '-', '-', '-', '-', '-'
            marks_lst = []

            for j in y:
                name_rollno = [roll, name[1:]]
                if y[0].isdigit() and j.isalpha() and j != ('PASS' or 'FAIL'):
                    name = name + j + ' '
                elif y[0].isdigit() and j.isdigit() and len(j) > 4:
                    roll += j

            if len(name_rollno[0]) > 1 and len(name_rollno[1]) > 1:
                lst.append(name_rollno)

            if y[0].isdigit() and len(y[0]) > 4:
                sub_codes = []
                marks_lst = [english, maths, phy, che, cs, ap_mat, bio, eco, acc, bs]

                for code in y:
                    if code.isdigit() and len(code) == 3:
                        sub_codes.append(code)

                y_1 = x[i + 1].split()
                for pos in range(len(y_1)):
                    try:
                        if y_1[pos].isdigit() == False:
                            y_1.remove(y_1[pos])
                    except:
                        break

                for m in range(len(sub_codes)):
                    if sub_codes[m] == '301':
                        marks_lst[0] = y_1[m]
                    elif sub_codes[m] == '041':
                        marks_lst[1] = y_1[m]
                    elif sub_codes[m] == '042':
                        marks_lst[2] = y_1[m]
                    elif sub_codes[m] == '043':
                        marks_lst[3] = y_1[m]
                    elif sub_codes[m] == '083':
                        marks_lst[4] = y_1[m]
                    elif sub_codes[m] == '241':
                        marks_lst[5] = y_1[m]
                    elif sub_codes[m] == '044':
                        marks_lst[6] = y_1[m]
                    elif sub_codes[m] == '030':
                        marks_lst[7] = y_1[m]
                    elif sub_codes[m] == '055':
                        marks_lst[8] = y_1[m]
                    elif sub_codes[m] == '054':
                        marks_lst[9] = y_1[m]
                    else:
                        print('New sub_code found!', y[m])

            if marks_lst == []:
                pass
            else:
                l2.append(marks_lst)

    for j in range(0, len(lst)):
        lk = lst[j] + l2[j]
        final.append(lk)

    if opt=='':
        eng = []
        mat = []
        physics = []
        chemistry = []
        C_S = []
        ap_maths = []
        biology = []
        economics = []
        account = []
        b_s = []
        for a in final:
            if a[2].isdigit():
                eng.append(a[1:3])
            if a[3].isdigit():
                mat.append([a[1]] + [a[3]])
            if a[4].isdigit():
                physics.append([a[1]] + [a[4]])
            if a[5].isdigit():
                chemistry.append([a[1]] + [a[5]])
            if a[6].isdigit():
                C_S.append([a[1]] + [a[6]])
            if a[7].isdigit():
                ap_maths.append([a[1]] + [a[7]])
            if a[8].isdigit():
                biology.append([a[1]] + [a[8]])
            if a[9].isdigit():
                economics.append([a[1]] + [a[9]])
            if a[10].isdigit():
                account.append([a[1]] + [a[10]])
            if a[11].isdigit():
                b_s.append([a[1]] + [a[11]])
        sorting(eng); sorting(mat); sorting(physics);  sorting(chemistry);  sorting(C_S);  sorting(ap_maths);  sorting(biology);  sorting(economics);  sorting(account);  sorting(b_s)

        count_eng = 1
        for t in range(len(eng)):
            if count_eng == 4:
                eng = eng[:t]
                break
            elif eng[t][-1] != eng[t + 1][-1]:
                eng[t].insert(0, count_eng)
                count_eng += 1
            else:
                eng[t].insert(0, count_eng)

        count_mat = 1
        for t in range(len(mat)):
            if count_mat == 4:
                mat = mat[:t]
                break
            elif mat[t][-1] != mat[t + 1][-1]:
                mat[t].insert(0, count_mat)
                count_mat += 1
            else:
                mat[t].insert(0, count_mat)

        count_phy = 1
        for t in range(len(physics)):
            if count_phy == 4:
                physics = physics[:t]
                break
            elif physics[t][-1] != physics[t + 1][-1]:
                physics[t].insert(0, count_phy)
                count_phy += 1
            else:
                physics[t].insert(0, count_phy)

        count_che = 1
        for t in range(len(chemistry)):
            if count_che == 4:
                chemistry = chemistry[:t]
                break
            elif chemistry[t][-1] != chemistry[t + 1][-1]:
                chemistry[t].insert(0, count_che)
                count_che += 1
            else:
                chemistry[t].insert(0, count_che)

        count_cs = 1
        for t in range(len(C_S)):
            if count_cs == 4:
                C_S = C_S[:t]
                break
            elif C_S[t][-1] != C_S[t + 1][-1]:
                C_S[t].insert(0, count_cs)
                count_cs += 1
            else:
                C_S[t].insert(0, count_cs)

        count_ap_mat = 1
        for t in range(len(ap_maths)):
            if count_ap_mat == 4:
                ap_maths = ap_maths[:t]
                break
            elif ap_maths[t][-1] != ap_maths[t + 1][-1]:
                ap_maths[t].insert(0, count_ap_mat)
                count_ap_mat += 1
            else:
                ap_maths[t].insert(0, count_ap_mat)

        count_bio = 1
        for t in range(len(biology)):
            if count_bio == 4:
                biology = biology[:t]
                break
            elif biology[t][-1] != biology[t + 1][-1]:
                biology[t].insert(0, count_bio)
                count_bio += 1
            else:
                biology[t].insert(0, count_bio)

        count_eco = 1
        for t in range(len(economics)):
            if count_eco == 4:
                economics = economics[:t]
                break
            elif economics[t][-1] != economics[t + 1][-1]:
                economics[t].insert(0, count_eco)
                count_eco += 1
            else:
                economics[t].insert(0, count_eco)

        count_acc = 1
        for t in range(len(account)):
            if count_acc == 4:
                account = account[:t]
                break
            elif account[t][-1] != account[t + 1][-1]:
                account[t].insert(0, count_acc)
                count_acc += 1
            else:
                account[t].insert(0, count_acc)

        count_bs = 1
        for t in range(len(b_s)):
            if count_bs == 4:
                b_s = b_s[:t]
                break
            elif b_s[t][-1] != b_s[t + 1][-1]:
                b_s[t].insert(0, count_bs)
                count_bs += 1
            else:
                b_s[t].insert(0, count_bs)

        def table(sub):
            if sub == eng:
                doc.add_heading("English Toppers", 0)
            if sub == mat:
                doc.add_heading("Maths Toppers", 0)
            if sub == physics:
                doc.add_heading("Physics Toppers", 0)
            if sub == chemistry:
                doc.add_heading("Chemistry Toppers", 0)
            if sub == C_S:
                doc.add_heading("C.S Toppers", 0)
            if sub == ap_maths:
                doc.add_heading("App.Maths Toppers", 0)
            if sub == biology:
                doc.add_heading("Biology Toppers", 0)
            if sub == economics:
                doc.add_heading("Economics Toppers", 0)
            if sub == account:
                doc.add_heading("Accountancy Toppers", 0)
            if sub == b_s:
                doc.add_heading("B.S Toppers", 0)

            table = doc.add_table(rows=1, cols=3)
            table.autofit = True
            row = table.rows[0].cells
            row[0].text = 'Place'
            row[1].text = 'Name'
            row[2].text = 'Mark'

            for place, name, mark in sub:
                row = table.add_row().cells
                row[0].text = str(place)
                row[1].text = name
                row[2].text = str(int(mark))
            doc.add_page_break()
            doc.save(file[:-4] + '.docx')

        table(eng); table(mat); table(physics); table(chemistry); table(C_S); table(ap_maths); table(biology); table(economics); table(account); table(b_s)

    else:
        name_lst = []
        for i in x1:
            name_lst.append(i.strip())

        ds_lst = []
        bd_lst = []
        for i in final:
            if i[0] not in name_lst:
                ds_lst.append(i)
            elif i[0] in name_lst:
                bd_lst.append(i)

        def option(var, f_ext):
            eng = []
            mat = []
            physics = []
            chemistry = []
            C_S = []
            ap_maths = []
            biology = []
            economics = []
            account = []
            b_s = []
            for a in var:
                if a[2].isdigit():
                    eng.append(a[1:3])
                if a[3].isdigit():
                    mat.append([a[1]] + [a[3]])
                if a[4].isdigit():
                    physics.append([a[1]] + [a[4]])
                if a[5].isdigit():
                    chemistry.append([a[1]] + [a[5]])
                if a[6].isdigit():
                    C_S.append([a[1]] + [a[6]])
                if a[7].isdigit():
                    ap_maths.append([a[1]] + [a[7]])
                if a[8].isdigit():
                    biology.append([a[1]] + [a[8]])
                if a[9].isdigit():
                    economics.append([a[1]] + [a[9]])
                if a[10].isdigit():
                    account.append([a[1]] + [a[10]])
                if a[11].isdigit():
                    b_s.append([a[1]] + [a[11]])
            sorting(eng); sorting(mat); sorting(physics); sorting(chemistry); sorting(C_S); sorting(ap_maths); sorting(biology); sorting(economics); sorting(account); sorting(b_s)

            count_eng = 1
            if len(eng) != 0:
                for t in range(len(eng)):
                    if t == len(eng) - 1:
                        if eng[t][-1] != eng[t - 1][-1]:
                            eng[t].insert(0, count_eng)
                        elif count_eng != 1:
                            eng[t].insert(0, count_eng - 1)
                        else:
                            eng[t].insert(0, count_eng)
                    elif count_eng == 4:
                        eng = eng[:t]
                        break
                    elif eng[t][-1] != eng[t + 1][-1]:
                        eng[t].insert(0, count_eng)
                        count_eng += 1
                    else:
                        eng[t].insert(0, count_eng)
            else:
                pass

            count_mat = 1
            for t in range(len(mat)):
                if count_mat == 4:
                    mat = mat[:t]
                    break
                elif mat[t][-1] != mat[t + 1][-1]:
                    mat[t].insert(0, count_mat)
                    count_mat += 1
                else:
                    mat[t].insert(0, count_mat)

            count_phy = 1
            if len(physics) != 0:
                for t in range(len(physics)):
                    if t == len(physics) - 1:
                        if physics[t][-1] != physics[t - 1][-1]:
                            physics[t].insert(0, count_phy)
                        elif count_phy != 1:
                            physics[t].insert(0, count_phy - 1)
                        else:
                            physics[t].insert(0, count_phy)
                    elif count_phy == 4:
                        physics = physics[:t]
                        break
                    elif physics[t][-1] != physics[t + 1][-1]:
                        physics[t].insert(0, count_phy)
                        count_phy += 1
                    else:
                        physics[t].insert(0, count_phy)
            else:
                pass

            count_che = 1
            if len(chemistry) != 0:
                for t in range(len(chemistry)):
                    if t == len(chemistry) - 1:
                        if chemistry[t][-1] != chemistry[t - 1][-1]:
                            chemistry[t].insert(0, count_che)
                        elif count_che != 1:
                            chemistry[t].insert(0, count_che - 1)
                        else:
                            chemistry[t].insert(0, count_che)
                    elif count_che == 4:
                        chemistry = chemistry[:t]
                        break
                    elif chemistry[t][-1] != chemistry[t + 1][-1]:
                        chemistry[t].insert(0, count_che)
                        count_che += 1
                    else:
                        chemistry[t].insert(0, count_che)
            else:
                pass

            count_cs = 1
            if len(C_S) != 0:
                for t in range(len(C_S)):
                    if count_cs == 4:
                        C_S = C_S[:t]
                        break
                    elif t == len(C_S) - 1:
                        if C_S[t][-1] != C_S[t - 1][-1]:
                            C_S[t].insert(0, count_cs)
                        elif count_cs != 1:
                            C_S[t].insert(0, count_cs - 1)
                        else:
                            C_S[t].insert(0, count_cs)
                    elif C_S[t][-1] != C_S[t + 1][-1]:
                        C_S[t].insert(0, count_cs)
                        count_cs += 1
                    else:
                        C_S[t].insert(0, count_cs)
            else:
                pass

            count_ap_mat = 1
            if len(ap_maths) != 0:
                for t in range(len(ap_maths)):
                    if t == len(ap_maths) - 1:
                        if ap_maths[t][-1] != ap_maths[t - 1][-1]:
                            ap_maths[t].insert(0, count_ap_mat)
                        elif count_ap_mat != 1:
                            ap_maths[t].insert(0, count_ap_mat - 1)
                        else:
                            ap_maths[t].insert(0, count_ap_mat)
                    elif count_ap_mat == 4:
                        ap_maths = ap_maths[:t]
                        break
                    elif ap_maths[t][-1] != ap_maths[t + 1][-1]:
                        ap_maths[t].insert(0, count_ap_mat)
                        count_ap_mat += 1
                    else:
                        ap_maths[t].insert(0, count_ap_mat)
            else:
                pass

            count_bio = 1
            if len(biology) != 0:
                for t in range(len(biology)):
                    if t == len(biology) - 1:
                        if biology[t][-1] != biology[t - 1][-1]:
                            biology[t].insert(0, count_bio)
                        elif count_bio != 1:
                            biology[t].insert(0, count_bio - 1)
                        else:
                            biology[t].insert(0, count_bio)
                    elif count_bio == 4:
                        biology = biology[:t]
                        break
                    elif biology[t][-1] != biology[t + 1][-1]:
                        biology[t].insert(0, count_bio)
                        count_bio += 1
                    else:
                        biology[t].insert(0, count_bio)
            else:
                pass

            count_eco = 1
            if len(economics) != 0:
                for t in range(len(economics)):
                    if t == len(economics) - 1:
                        if economics[t][-1] != economics[t - 1][-1]:
                            economics[t].insert(0, count_eco)
                        elif count_eco != 1:
                            economics[t].insert(0, count_eco - 1)
                        else:
                            economics[t].insert(0, count_eco)
                    elif count_eco == 4:
                        economics = economics[:t]
                        break
                    elif economics[t][-1] != economics[t + 1][-1]:
                        economics[t].insert(0, count_eco)
                        count_eco += 1
                    else:
                        economics[t].insert(0, count_eco)
            else:
                pass

            count_acc = 1
            if len(account) != 0:
                for t in range(len(account)):
                    if t == len(account) - 1:
                        if account[t][-1] != account[t - 1][-1]:
                            account[t].insert(0, count_acc)
                        elif count_acc != 1:
                            account[t].insert(0, count_acc - 1)
                        else:
                            account[t].insert(0, count_acc)
                    elif count_acc == 4:
                        account = account[:t]
                        break
                    elif account[t][-1] != account[t + 1][-1]:
                        account[t].insert(0, count_acc)
                        count_acc += 1
                    else:
                        account[t].insert(0, count_acc)
            else:
                pass

            count_bs = 1
            if len(b_s) != 0:
                for t in range(len(b_s)):
                    if t == len(b_s) - 1:
                        if b_s[t][-1] != b_s[t - 1][-1]:
                            b_s[t].insert(0, count_bs)
                        elif count_bs != 1:
                            b_s[t].insert(0, count_bs - 1)
                        else:
                            b_s[t].insert(0, count_bs)
                    elif count_bs == 4:
                        b_s = b_s[:t]
                        break
                    elif b_s[t][-1] != b_s[t + 1][-1]:
                        b_s[t].insert(0, count_bs)
                        count_bs += 1
                    else:
                        b_s[t].insert(0, count_bs)
            else:
                pass

            def table(sub,f_ext):
                if sub == eng:
                    f_ext.add_heading("English Toppers", 0)
                if sub == mat:
                    f_ext.add_heading("Maths Toppers", 0)
                if sub == physics:
                    f_ext.add_heading("Physics Toppers", 0)
                if sub == chemistry:
                    f_ext.add_heading("Chemistry Toppers", 0)
                if sub == C_S:
                    f_ext.add_heading("C.S Toppers", 0)
                if sub == ap_maths:
                    f_ext.add_heading("App.Maths Toppers", 0)
                if sub == biology:
                    f_ext.add_heading("Biology Toppers", 0)
                if sub == economics:
                    f_ext.add_heading("Economics Toppers", 0)
                if sub == account:
                    f_ext.add_heading("Accountancy Toppers", 0)
                if sub == b_s:
                    f_ext.add_heading("B.S Toppers", 0)

                table = f_ext.add_table(rows=1, cols=3)
                row = table.rows[0].cells
                row[0].text = 'Place'
                row[1].text = 'Name'
                row[2].text = 'Mark'

                for place, name, mark in sub:
                    row = table.add_row().cells
                    row[0].text = str(place)
                    row[1].text = name
                    row[2].text = str(int(mark))
                    f_ext.add_page_break()
                    if f_ext == doc:
                        f_ext.save(file[:-4] + '(DS).docx')
                    else:
                        f_ext.save(file[:-4] + '(BD).docx')

            if len(eng) != 0:
                table(eng, f_ext)
            if len(mat) != 0:
                table(mat, f_ext)
            if len(physics) != 0:
                table(physics,f_ext)
            if len(chemistry) != 0:
                table(chemistry,f_ext)
            if len(C_S) != 0:
                table(C_S,f_ext)
            if len(ap_maths) != 0:
                table(ap_maths,f_ext)
            if len(biology) != 0:
                table(biology,f_ext)
            if len(economics) != 0:
                table(economics,f_ext)
            if len(account) != 0:
                table(account,f_ext)
            if len(b_s) != 0:
                table(b_s,f_ext)

            eng.clear(); mat.clear(); physics.clear(); chemistry.clear(); C_S.clear(); ap_maths.clear(); biology.clear(); economics.clear(); account.clear(); b_s.clear()
        print(len(ds_lst))
        print(len(bd_lst))
        option(ds_lst, doc)
        option(bd_lst, doc1)
